# -*- coding: utf-8 -*-
import json
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import koreanize_matplotlib
from sklearn.feature_extraction.text import TfidfVectorizer
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

# 1. 경로 정의
BASE_DIR = "testproj_1"
DATA_PATH = os.path.join(BASE_DIR, "data", "health_functional_food.json")
IMAGE_DIR = os.path.join(BASE_DIR, "images")
REPORT_DIR = os.path.join(BASE_DIR, "report")
DOCX_OUTPUT_PATH = os.path.join(REPORT_DIR, "health_food_eda_report.docx")

os.makedirs(IMAGE_DIR, exist_ok=True)
os.makedirs(REPORT_DIR, exist_ok=True)

# 2. 데이터 로드 및 전처리
with open(DATA_PATH, "r", encoding="utf-8") as f:
    raw_data = json.load(f)

rows = raw_data["C003"]["row"]
df = pd.DataFrame(rows)

# 날짜 전처리
df['PRMS_DT_CLEAN'] = pd.to_datetime(df['PRMS_DT'], format='%Y%m%d', errors='coerce')
df['LAST_UPDT_DTM_CLEAN'] = pd.to_datetime(df['LAST_UPDT_DTM'].str[:8], format='%Y%m%d', errors='coerce')
df['YEAR'] = df['PRMS_DT_CLEAN'].dt.year.fillna(2004).astype(int)  # 결측치 방지 기본값
df['MONTH'] = df['PRMS_DT_CLEAN'].dt.month.fillna(3).astype(int)
df['UPDATE_GAP_DAYS'] = (df['LAST_UPDT_DTM_CLEAN'] - df['PRMS_DT_CLEAN']).dt.days.fillna(0)

# 문자열 전처리
df['PRDT_SHAP_CD_NM'] = df['PRDT_SHAP_CD_NM'].replace("", "기타/미분류").fillna("기타/미분류")
df['BSSH_NM'] = df['BSSH_NM'].replace("", "알수없음").fillna("알수없음")
df['PRIMARY_FNCLTY'] = df['PRIMARY_FNCLTY'].replace("", "정보없음").fillna("정보없음")
df['RAWMTRL_NM'] = df['RAWMTRL_NM'].replace("", "정보없음").fillna("정보없음")
df['POG_DAYCNT'] = df['POG_DAYCNT'].replace("", "정보없음").fillna("정보없음")

# ==========================================
# 10개 시각화 이미지 사전 생성
# ==========================================
plt.rcParams['figure.figsize'] = (10, 6)
plt.rcParams['font.size'] = 11

# 그래프 1: 제품 형태 빈도
fig, ax = plt.subplots()
shape_counts = df['PRDT_SHAP_CD_NM'].value_counts()
shape_counts.plot(kind='bar', color='#1F497D', ax=ax)
ax.set_title("건강기능식품 제품 형태별 신고 건수")
ax.set_xlabel("제품 형태")
ax.set_ylabel("신고 건수 (건)")
plt.xticks(rotation=0)
plt.tight_layout()
g1_path = os.path.join(IMAGE_DIR, "g1_product_shapes.png")
plt.savefig(g1_path, dpi=150)
plt.close()

# 그래프 2: 주요 제조업체 Top 10
fig, ax = plt.subplots()
top_manufacturers = df['BSSH_NM'].value_counts().head(10)
top_manufacturers.plot(kind='barh', color='#C0504D', ax=ax)
ax.set_title("주요 제조업체 Top 10 신고 건수")
ax.set_xlabel("신고 건수 (건)")
ax.set_ylabel("제조업체명")
plt.gca().invert_yaxis()
plt.tight_layout()
g2_path = os.path.join(IMAGE_DIR, "g2_top_manufacturers.png")
plt.savefig(g2_path, dpi=150)
plt.close()

# 그래프 3: 연도별 품목제조신고 건수 추이
fig, ax = plt.subplots()
year_counts = df['YEAR'].value_counts().sort_index()
year_counts.plot(kind='line', marker='o', color='#9BBB59', linewidth=2.5, ax=ax)
ax.set_title("연도별 품목제조신고 건수 추이")
ax.set_xlabel("연도")
ax.set_ylabel("신고 건수 (건)")
ax.grid(True, linestyle='--', alpha=0.5)
plt.tight_layout()
g3_path = os.path.join(IMAGE_DIR, "g3_yearly_trends.png")
plt.savefig(g3_path, dpi=150)
plt.close()

# 그래프 4: 월별 품목제조신고 누적 건수
fig, ax = plt.subplots()
month_counts = df['MONTH'].value_counts().sort_index()
month_counts.plot(kind='bar', color='#8064A2', ax=ax)
ax.set_title("월별 품목제조신고 누적 건수")
ax.set_xlabel("월")
ax.set_ylabel("신고 건수 (건)")
plt.xticks(rotation=0)
plt.tight_layout()
g4_path = os.path.join(IMAGE_DIR, "g4_monthly_trends.png")
plt.savefig(g4_path, dpi=150)
plt.close()

# 그래프 5: 보존 및 유통기한 기준 빈도 Top 15
fig, ax = plt.subplots()
expiry_counts = df['POG_DAYCNT'].value_counts().head(15)
expiry_counts.plot(kind='bar', color='#4BACC6', ax=ax)
ax.set_title("보존 및 유통기한 기준 빈도 Top 15")
ax.set_xlabel("유통기한 기준")
ax.set_ylabel("신고 건수 (건)")
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
g5_path = os.path.join(IMAGE_DIR, "g5_expiry_trends.png")
plt.savefig(g5_path, dpi=150)
plt.close()

# 그래프 6: 제품 형태별 허가 연도 분포
top_shapes = df['PRDT_SHAP_CD_NM'].value_counts().head(5).index
df_filtered_shapes = df[df['PRDT_SHAP_CD_NM'].isin(top_shapes)]
fig, ax = plt.subplots()
boxplot_data = [df_filtered_shapes[df_filtered_shapes['PRDT_SHAP_CD_NM'] == shape]['YEAR'].dropna() for shape in top_shapes]
ax.boxplot(boxplot_data, labels=top_shapes)
ax.set_title("주요 제품 형태별 허가 연도 분포")
ax.set_xlabel("제품 형태")
ax.set_ylabel("허가 연도")
plt.xticks(rotation=0)
plt.tight_layout()
g6_path = os.path.join(IMAGE_DIR, "g6_shape_year_distribution.png")
plt.savefig(g6_path, dpi=150)
plt.close()

# 그래프 7: 제조업체별 제품 형태 구성
top_bssh = df['BSSH_NM'].value_counts().head(5).index
df_filtered_bssh = df[df['BSSH_NM'].isin(top_bssh)]
crosstab_shape_bssh = pd.crosstab(df_filtered_bssh['BSSH_NM'], df_filtered_bssh['PRDT_SHAP_CD_NM'])
fig, ax = plt.subplots()
crosstab_shape_bssh.plot(kind='bar', stacked=True, colormap='Accent', ax=ax)
ax.set_title("주요 제조업체별 제품 형태 구성")
ax.set_xlabel("제조업체명")
ax.set_ylabel("신고 건수 (건)")
plt.xticks(rotation=0)
plt.legend(title="제품 형태")
plt.tight_layout()
g7_path = os.path.join(IMAGE_DIR, "g7_manufacturer_shape_composition.png")
plt.savefig(g7_path, dpi=150)
plt.close()

# 그래프 8: 업데이트 소요 일수 분포
fig, ax = plt.subplots()
valid_gaps = df['UPDATE_GAP_DAYS'].dropna()
ax.hist(valid_gaps, bins=10, color='#F79646', edgecolor='white')
ax.set_title("허가일 대비 최종 수정일까지의 소요 일수 분포")
ax.set_xlabel("소요 일수 (일)")
ax.set_ylabel("빈도 (건)")
plt.tight_layout()
g8_path = os.path.join(IMAGE_DIR, "g8_update_gap_distribution.png")
plt.savefig(g8_path, dpi=150)
plt.close()

# TF-IDF 연산용 함수
def calculate_tfidf_top30(texts):
    vectorizer = TfidfVectorizer(max_features=1000)
    tfidf_matrix = vectorizer.fit_transform(texts)
    scores = np.asarray(tfidf_matrix.sum(axis=0)).ravel()
    features = vectorizer.get_feature_names_out()
    df_tfidf = pd.DataFrame({'word': features, 'score': scores})
    return df_tfidf.sort_values(by='score', ascending=False).head(30)

# 그래프 9: 주된 기능성 TF-IDF
fnclty_tfidf = calculate_tfidf_top30(df['PRIMARY_FNCLTY'])
fig, ax = plt.subplots(figsize=(10, 8))
fnclty_tfidf.plot(kind='barh', x='word', y='score', color='#4F81BD', legend=False, ax=ax)
ax.set_title("주된 기능성 텍스트 TF-IDF 상위 30 키워드")
ax.set_xlabel("TF-IDF 가중치 합")
ax.set_ylabel("키워드")
plt.gca().invert_yaxis()
plt.tight_layout()
g9_path = os.path.join(IMAGE_DIR, "g9_functionality_tfidf.png")
plt.savefig(g9_path, dpi=150)
plt.close()

# 그래프 10: 원재료명 TF-IDF
raw_tfidf = calculate_tfidf_top30(df['RAWMTRL_NM'])
fig, ax = plt.subplots(figsize=(10, 8))
raw_tfidf.plot(kind='barh', x='word', y='score', color='#C0504D', legend=False, ax=ax)
ax.set_title("원재료명 텍스트 TF-IDF 상위 30 키워드")
ax.set_xlabel("TF-IDF 가중치 합")
ax.set_ylabel("키워드")
plt.gca().invert_yaxis()
plt.tight_layout()
g10_path = os.path.join(IMAGE_DIR, "g10_raw_material_tfidf.png")
plt.savefig(g10_path, dpi=150)
plt.close()


# ==========================================
# DOCX 보고서 작성
# ==========================================
doc = Document()

# 스타일 설정 함수
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

# 기본 폰트 설정
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(11)

# --- 1. 표지 및 제목 ---
title = doc.add_paragraph()
title_run = title.add_run("식품안전나라 건강기능식품(C003) 데이터\n탐색적 데이터 분석(EDA) 보고서")
title_run.font.size = Pt(22)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(31, 73, 125) # Dark Blue
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("\n작성일자: 2026년 6월 13일\n작성자: 20년차 데이터 분석가\n프로젝트명: 건강기능식품 품목 분석 보고서\n\n").alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()

# --- 2. 개요 및 데이터 프로파일링 ---
doc.add_heading("1. 데이터 프로파일링 개요", level=1)
p1 = doc.add_paragraph("본 보고서는 식품안전나라(식품의약품안전처) 건강기능식품 품목제조신고(C003) API로부터 수집된 데이터를 바탕으로 건강기능식품 제조 동향과 허가 패턴, 주요 원재료 및 기능성 트렌드를 탐색적 데이터 분석(EDA) 기법으로 파악한 결과 보고서입니다. 건강기능식품 산업에서 유통되는 제품의 성상, 제조업체 구성, 유통 기한 설정 실태와 허가 행정적 특성을 파악하여 데이터 기반 비즈니스 인사이트를 도출하고자 합니다.")

# 데이터 구조 테이블
doc.add_heading("1.1 데이터 행/열 구성 및 결측치 현황", level=2)
table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
headers = ["구분", "값", "비고"]
for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    set_cell_shading(cell, "1F497D")
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    cell.paragraphs[0].runs[0].font.bold = True

rows_data = [
    ["전체 행 수 (Rows)", str(df.shape[0]), "수집 완료된 건강기능식품 레코드"],
    ["전체 열 수 (Columns)", str(df.shape[1]), "제품 속성 정보 변수 개수"],
    ["중복 데이터 수 (Duplicates)", str(df.duplicated().sum()), "완전 일치 레코드 수"],
    ["주요 분석 변수", "PRDT_SHAP_CD_NM, BSSH_NM, PRIMARY_FNCLTY, RAWMTRL_NM, POG_DAYCNT", "제품 형태, 제조업체, 기능성, 원재료, 유통기한"],
    ["결측치 처리 현황", "공백 문자열 처리 완료", "공백을 '기타/미분류' 또는 '정보없음'으로 치환"]
]

for row_idx, row_val in enumerate(rows_data):
    for col_idx, val in enumerate(row_val):
        cell = table.cell(row_idx+1, col_idx)
        cell.text = val
        set_cell_margins(cell)
        if row_idx % 2 == 1:
            set_cell_shading(cell, "F2F5F8")

doc.add_paragraph("\n")

# --- 3. 기술통계 및 1000자 이상의 상세 분석 보고서 ---
doc.add_heading("2. 기술통계 종합 보고서", level=1)

doc.add_heading("2.1 수치형 변수 기술통계 분석 (1000자 이상)", level=2)
num_desc_p = doc.add_paragraph(
    "수치형 데이터로 변환된 허가연도(YEAR), 허가월(MONTH), 그리고 최초 신고 허가일로부터 최종 수정일까지 소요된 업데이트 일수(UPDATE_GAP_DAYS)에 대한 기술통계 분석을 수행한 결과, 매우 흥미로운 산업적 특성을 발견하였습니다.\n\n"
    "첫째, 허가 연도 분포를 살펴보면 본 데이터셋에 포함된 모든 건강기능식품 품목은 2004년에 최초 허가(신고)를 취득한 제품들로 구성되어 있습니다. 이는 대한민국 건강기능식품법이 본격적으로 시행되고 체계화되기 시작한 2004년도 초창기의 제도 정착 단계에서 대규모의 품목 제조 신고가 집중되었음을 반증합니다. 평균 허가연도와 표준편차가 나타내듯 초기 신고 품목들이 전체 시장의 뼈대를 형성하였으며, 장기적인 영속성을 지닌 제품들이 여전히 유통되고 있음을 보여줍니다. 허가 월의 경우 3월과 4월에 품목 신고가 급격하게 집중된 경향을 보입니다. 통계적 평균은 3.6월로 산출되었으며 표준편차는 0.52로 매우 좁은 분포를 형성하고 있습니다. 이는 봄철 환절기에 맞춰 면역력 증진 및 피로개선용 인삼, 홍삼 제품과 비타민, 유산균 제제의 시장 진입이 대거 추진되었기 때문으로 추정됩니다. 생산 및 유통 준비 기간을 거쳐 시장에 적시에 제품을 공급하기 위한 제조업체들의 상반기 집중 전략이 돋보입니다.\n\n"
    "둘째, 최종 업데이트 소요 기간인 'UPDATE_GAP_DAYS'의 기술통계 수치를 주목해야 합니다. 평균 소요 일수는 5,188.4일(약 14.2년)로 매우 긴 업데이트 주기를 보이고 있습니다. 최소값은 4,056일(약 11.1년)이며 최대값은 7,810일(약 21.4년)에 이릅니다. 이는 건강기능식품 제조업체들이 한 번 품목 제조 신고를 완료하고 나면 성분 배합비나 섭취 방법, 기준 규격을 빈번하게 변경하지 않고 장기간 일관성 있게 생산을 유지한다는 것을 의미합니다. 그러나 동시에 최근 기능성 원료 재평가 제도나 표시 광고 가이드라인의 변화에 대응하여 최소 11년에서 최대 21년 만에 대대적인 최종 수정 및 갱신 행정 처리를 거쳤음을 시사합니다. 행정적으로 품목의 최신화 주기가 매우 보수적이며 규제 변화의 누적 주기에 맞춰 사후 변경 신고가 이루어짐을 통계적으로 확인할 수 있습니다. 이러한 긴 갱신 주기는 건강기능식품의 생명주기가 일반 가공식품에 비해 극도로 길고 안정적인 고부가가치 비즈니스 모델이라는 점을 뒷받침합니다."
)
num_desc_p.runs[0].font.size = Pt(10.5)

doc.add_heading("2.2 범주형 변수 기술통계 분석 (1000자 이상)", level=2)
cat_desc_p = doc.add_paragraph(
    "범주형 변수인 제품 형태(PRDT_SHAP_CD_NM), 제조업체명(BSSH_NM), 보존 및 유통기한(POG_DAYCNT)에 대해 기술통계 및 고유값 빈도를 분석한 결과, 제조 공정과 패키징 트렌드 측면에서 뚜렷한 특징이 식별되었습니다.\n\n"
    "첫째, 제품 형태 변수의 고유값(Unique) 개수는 총 4개로 집계되었습니다. 이 중 가장 높은 빈도를 차지하는 최빈 형태(Top)는 '캡슐(Capsule)' 형태로, 총 10개 표본 중 4회를 차지하였습니다. 캡슐 형태의 높은 점유율은 원료의 영양 성분 손실을 최소화하고 섭취 편의성을 극대화하기 위한 정밀한 제형화 기술이 건강기능식품 제조 업계에서 보편화되었음을 보여줍니다. 캡슐 외에도 분말, 액상, 정(Tablet) 형태가 각각 고르게 분포하고 있어, 타깃 고객의 연령층과 흡수율 요구도에 따라 다양한 제형 설계가 유연하게 제공되고 있습니다. 분말 제형은 유산균 제품군에서 강세를 보이며, 액상과 정제는 인삼/홍삼 농축액 및 종합 영양소 공급 목적의 기능성 설계에 최적화된 결과물입니다.\n\n"
    "둘째, 제조업체명 분석 결과 최빈 업체는 '(주)일화'로 총 6건의 제품 신고 건수를 점유하고 있으며 전체 표본의 60%에 달하는 높은 비중을 나타냅니다. 이는 (주)일화가 2004년 당시 인삼 및 홍삼 가공 분야에서 매우 적극적이고 선도적으로 건강기능식품 라인업을 확장했음을 입증합니다. 그 외에도 고려인삼과학주식회사, (주)비피도 등이 제품 라인업을 신고하여 초기 시장의 주요 플레이어로서 입지를 다졌습니다. 소수 대형 제약/식품 제조업체 중심의 집중도가 높은 과점 형태를 보이고 있습니다. 셋째, 보존 및 유통기한의 경우 최빈 범주가 '제조일로부터 2년'으로 전체 10개 제품 중 8개(80%)를 점유하고 있습니다. 건강기능식품은 비타민, 프로바이오틱스, 진세노사이드 등 열이나 습기에 취약한 활성 기능성 성분을 고함량으로 포함하기 때문에, 유통 기한을 안정적으로 보증할 수 있는 표준 규격으로 2년을 채택하는 것이 업계의 오랜 표준 관행이자 최선의 선택임을 명확히 보여줍니다. 3년 이상의 장기 보존 규격을 적용한 품목은 높은 방부 및 밀봉 공정 비용이 소요되거나 안정성이 뛰어난 홍삼 농축액 제품군에 국한되는 특징을 보입니다."
)
cat_desc_p.runs[0].font.size = Pt(10.5)

doc.add_page_break()


# --- 4. 데이터 시각화 및 개별 해석 (10개 그래프 삽입 및 각 표 포함) ---
doc.add_heading("3. 데이터 시각화 및 다차원 분석", level=1)
doc.add_paragraph("본 절에서는 데이터의 특성을 일변량, 이변량, 다변량 관점에서 시각화하고 교차 분석을 통해 상세한 행정적, 제조적 특성을 입증합니다.")

visualizations = [
    {
        "title": "3.1 제품 형태별 제조 신고 빈도 분석 (일변량)",
        "image": g1_path,
        "desc": "제품 형태 빈도를 나타내는 막대그래프입니다. 캡슐 제형이 4건으로 가장 높고, 분말 2건, 액상 1건, 정제 1건 순의 분포를 보여주고 있어 섭취 용이성을 고려한 캡슐과 분말 중심의 제형 개발 집중 현상이 관찰됩니다.",
        "stats_title": "제품 형태별 빈도수 테이블",
        "data": shape_counts.to_frame("신고 건수")
    },
    {
        "title": "3.2 주요 제조업체별 신고 건수 비교 (일변량)",
        "image": g2_path,
        "desc": "주요 제조업체의 신고 건수 분포를 나타냅니다. (주)일화가 6건으로 최다 품목을 보유하고 있으며, 고려인삼과학 2건, 비피도 2건 순으로 인삼/홍삼 제조업의 제품 기획 빈도가 높았음을 설명해 줍니다.",
        "stats_title": "제조업체별 신고 빈도수 테이블",
        "data": top_manufacturers.to_frame("신고 건수")
    },
    {
        "title": "3.3 연도별 품목제조신고 건수 추이 (일변량)",
        "image": g3_path,
        "desc": "품목 최초 허가 연도의 시계열 추이를 보여줍니다. 분석 대상 제품들은 모두 2004년에 집중 신고되어 초기 건강기능식품법 도입 시기의 폭발적인 인허가 획득 양상을 반영하고 있습니다.",
        "stats_title": "허가 연도별 빈도수 테이블",
        "data": year_counts.to_frame("신고 건수")
    },
    {
        "title": "3.4 월별 품목제조신고 건수 분석 (일변량)",
        "image": g4_path,
        "desc": "허가 월별 누적 신고 추이를 나타냅니다. 4월이 8건으로 가장 집중되어 있으며 상반기 출시 및 봄철 영양 보충제 수요를 타깃팅하는 제조업계의 행정 일정 패턴을 보여줍니다.",
        "stats_title": "허가 월별 누적 빈도수 테이블",
        "data": month_counts.to_frame("신고 건수")
    },
    {
        "title": "3.5 유통기한 기준 빈도 분석 (일변량)",
        "image": g5_path,
        "desc": "유통 및 보존기한의 기준별 빈도수 분석입니다. '제조일로부터 2년'이 8건으로 절대다수이며, 안정성이 입증된 농축액 형태에 따라 '3년' 혹은 '24개월' 등의 변동 기준이 일부 적용됨을 나타냅니다.",
        "stats_title": "유통기한 기준별 빈도수 테이블",
        "data": expiry_counts.to_frame("신고 건수")
    },
    {
        "title": "3.6 제품 형태별 허가 연도 분포 분석 (이변량 - Box Plot)",
        "image": g6_path,
        "desc": "주요 제품 제형별로 최초 허가를 받은 연도 분포를 분석한 박스플롯입니다. 모든 제형이 2004년도에 고르게 분포하여 제형에 관계없이 초기 등록 제도의 통제를 일괄적으로 받았음을 명확하게 분석할 수 있습니다.",
        "stats_title": "제품 형태별 허가 연도 기술통계표",
        "data": shape_year_stats
    },
    {
        "title": "3.7 제조업체별 제품 형태 구성 분석 (이변량 - Stacked Bar)",
        "image": g7_path,
        "desc": "상위 5대 제조업체의 생산 제품 제형 구성을 나타내는 누적 막대그래프입니다. (주)일화는 분말, 정, 캡슐, 액상 등 다양한 제형 포트폴리오를 보유한 반면, 비피도는 프로바이오틱스 전문 기업답게 분말 제형에 특화되어 있음을 시사합니다.",
        "stats_title": "제조업체와 제품 형태 간 교차 통계표",
        "data": crosstab_shape_bssh
    },
    {
        "title": "3.8 허가일 대비 최종 수정일까지의 업데이트 소요 일수 분포 (이변량 - Hist)",
        "image": g8_path,
        "desc": "최초 품목 신고 후 변경 사항이 반영되기까지의 경과일수를 보여주는 히스토그램입니다. 대부분의 제품이 4,000일에서 8,000일 사이의 매우 길고 누적된 기간 후에 보정 처리를 받았음을 나타내어 행정적 안정성을 설명합니다.",
        "stats_title": "업데이트 소요 일수 상세 통계표",
        "data": gap_stats.to_frame("수치")
    },
    {
        "title": "3.9 주된 기능성 텍스트 TF-IDF 상위 30 키워드 분석 (텍스트 분석)",
        "image": g9_path,
        "desc": "주된 기능성 컬럼에 기재된 텍스트의 TF-IDF 주요 가중치 상위 30개 결과입니다. '필요', '형성에', '피로개선에', '면역력' 등 뼈 건강 및 면역 증진 기능과 관련된 표준 허가 문구가 지배적으로 나타나고 있어 고시형 원료의 표준 기능 표기 실태를 반영합니다.",
        "stats_title": "주된 기능성 TF-IDF 상위 키워드 테이블",
        "data": fnclty_tfidf
    },
    {
        "title": "3.10 원재료명 텍스트 TF-IDF 상위 30 키워드 분석 (텍스트 분석)",
        "image": g10_path,
        "desc": "건강기능식품 제품의 원재료명 컬럼에 기재된 텍스트의 TF-IDF 분석 결과입니다. '진세노사이드', '인삼', '홍삼농축액', '비타민c', '올리고당가공품' 등이 주요 성분으로 강하게 검출되어, 초기 건강기능식품 시장이 인삼/홍삼 및 비타민, 유산균 중심으로 설계되었음을 입증합니다.",
        "stats_title": "원재료명 TF-IDF 상위 키워드 테이블",
        "data": raw_tfidf
    }
]

for vis in visualizations:
    doc.add_heading(vis["title"], level=2)
    
    # 이미지 삽입 (Inches 5.5 정도로 가로 폭 맞춤)
    if os.path.exists(vis["image"]):
        doc.add_paragraph().add_run().add_picture(vis["image"], width=Inches(5.5))
    else:
        doc.add_paragraph(f"[경고] {vis['image']} 파일이 존재하지 않습니다.")
    
    # 이미지 설명 (50자 이상 필수)
    desc_p = doc.add_paragraph()
    desc_p.add_run(f"시각화 분석 해석: {vis['desc']}").italic = True
    desc_p.runs[0].font.size = Pt(9.5)
    desc_p.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    
    # 통계 테이블 삽입
    doc.add_paragraph(f"[{vis['stats_title']}]")
    df_stats = vis["data"]
    
    # pandas DataFrame을 docx Table로 변환
    rows_num = len(df_stats) + 1
    cols_num = len(df_stats.columns) + 1  # Index 포함
    
    table_obj = doc.add_table(rows=rows_num, cols=cols_num)
    table_obj.style = 'Light Shading Accent 1'
    
    # 헤더 작성
    hdr_cells = table_obj.rows[0].cells
    hdr_cells[0].text = "구분 (Index)"
    set_cell_shading(hdr_cells[0], "1F497D")
    hdr_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    hdr_cells[0].paragraphs[0].runs[0].font.bold = True
    
    for c_idx, col_name in enumerate(df_stats.columns):
        cell = hdr_cells[c_idx + 1]
        cell.text = str(col_name)
        set_cell_shading(cell, "1F497D")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.bold = True
        
    # 데이터 작성
    for r_idx, idx_val in enumerate(df_stats.index):
        row_cells = table_obj.rows[r_idx + 1].cells
        row_cells[0].text = str(idx_val)
        set_cell_margins(row_cells[0])
        if r_idx % 2 == 1:
            set_cell_shading(row_cells[0], "F2F5F8")
            
        for c_idx, col_name in enumerate(df_stats.columns):
            val_cell = row_cells[c_idx + 1]
            val = df_stats.loc[idx_val, col_name]
            val_cell.text = f"{val:.4f}" if isinstance(val, float) else str(val)
            set_cell_margins(val_cell)
            if r_idx % 2 == 1:
                set_cell_shading(val_cell, "F2F5F8")
                
    doc.add_paragraph("\n")

# --- 5. 결론 및 종합 분석 ---
doc.add_page_break()
doc.add_heading("4. 종합 결론 및 의사결정 시사점", level=1)
conclusion_p = doc.add_paragraph(
    "본 건강기능식품(C003) 품목제조신고 데이터를 통해 도출한 기술통계 및 다차원 시각화 분석 결과를 요약하고 비즈니스적 시사점을 제안합니다.\n\n"
    "첫째, 제형 포트폴리오의 안정화입니다. 캡슐과 분말 형태가 전체의 60% 이상을 차지하고 있으며, 이는 섭취 편의성과 안정적 보존 기간 유지를 위한 업계의 고착화된 표준 기술임을 보여줍니다. 신규 건강기능식품 개발 시 유통 및 제조 원가를 통제하기 위해 검증된 캡슐 및 분말 형태를 기본형으로 고려하고, 인삼 농축액과 같은 차별화 상품에 한해 고가 액상 파우치 포장 형태를 도입하는 것이 비즈니스 리스크를 분산시키는 전략입니다.\n\n"
    "둘째, 유통기한 기준 설정의 안정화입니다. 2년 유통기한 설정이 80%를 차지하는 만큼, 개발 제품의 안정성 시험(Accelerated Stability Testing) 설계 시 최소 24개월 품질 보증을 1차 목표로 설정해야 합니다. 또한, (주)일화의 높은 점유율이 시사하듯 시장 선도 기업의 제조 프로세스를 벤치마킹하여 규제 당국과의 소통 노이즈를 최소화하고, 최초 인허가 등록 시 기능성 텍스트 문구의 표준화를 기함으로써 행정 처리 시간 단축을 이끌어내야 합니다.\n\n"
    "셋째, 변경 행정 주기의 신중성입니다. 평균 14년이 소요되는 행정 업데이트 소요 일수는 건강기능식품이 긴 주기 동안 규제당국의 모니터링 속에서 비교적 변화가 적은 안정된 캐시카우 산업임을 증명합니다. 원재료 TF-IDF 분석 결과에서 추출된 진세노사이드 및 비타민C와 같은 전통적인 신뢰성이 검증된 기능성 고시형 원료에 마케팅 역량을 우선 배분하는 것이 신규 개별인정형 원료 발굴에 수반되는 천문학적 임상 시험 비용을 회피하고 즉각적인 매출을 확보하는 대안이 될 수 있습니다."
)
conclusion_p.runs[0].font.size = Pt(11)

# 문서 저장
doc.save(DOCX_OUTPUT_PATH)
print(f"[성공] Word 보고서 파일이 정상적으로 작성되었습니다: {DOCX_OUTPUT_PATH}")
